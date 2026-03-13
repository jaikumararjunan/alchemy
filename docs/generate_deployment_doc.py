#!/usr/bin/env python3
"""
Generate docs/DEPLOYMENT.docx from docs/DEPLOYMENT.md.

Usage:
    pip install python-docx
    python docs/generate_deployment_doc.py
"""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    sys.exit("python-docx is required: pip install python-docx")


# ── Helpers ───────────────────────────────────────────────────────────────────


def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc):
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pb.append(bottom)
    pPr.append(pb)
    p.paragraph_format.space_after = Pt(0)
    return p


def apply_inline_code(para, text: str):
    """Add text to paragraph, rendering `backtick` spans in Courier."""
    parts = re.split(r"`([^`]+)`", text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        if i % 2 == 1:  # odd indices = backtick content
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC9, 0x2A, 0x2A)


# ── Main builder ──────────────────────────────────────────────────────────────


def build_docx(md_path: Path, out_path: Path):
    md = md_path.read_text(encoding="utf-8")
    lines = md.splitlines()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Default body style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    in_code = False
    code_lines: list[str] = []
    in_table = False
    table_rows: list[list[str]] = []

    def flush_code():
        nonlocal code_lines
        if not code_lines:
            return
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run("\n".join(code_lines))
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        # Light grey shading
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F5F5F5")
        pPr.append(shd)
        para.paragraph_format.left_indent = Inches(0.3)
        code_lines = []

    def flush_table():
        nonlocal table_rows
        if not table_rows:
            return
        # Filter out separator rows (---|---|--)
        data_rows = [
            r
            for r in table_rows
            if not all(re.match(r"^[-: ]+$", c.strip()) for c in r)
        ]
        if not data_rows:
            table_rows = []
            return
        cols = max(len(r) for r in data_rows)
        tbl = doc.add_table(rows=len(data_rows), cols=cols)
        tbl.style = "Table Grid"
        for ri, row in enumerate(data_rows):
            for ci, cell_text in enumerate(row):
                if ci >= cols:
                    break
                cell = tbl.cell(ri, ci)
                cell.text = cell_text.strip()
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                if ri == 0:
                    cell.paragraphs[0].runs[0].bold = True
                    set_cell_bg(cell, "D9E8F5")
        doc.add_paragraph()  # spacing
        table_rows = []

    while i < len(lines):
        line = lines[i]

        # ── Code fences ──────────────────────────────────────────────────────
        if line.strip().startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
                in_table = False
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── Tables ───────────────────────────────────────────────────────────
        if line.startswith("|"):
            in_table = True
            cells = [c for c in line.split("|") if c != ""]
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            flush_table()
            in_table = False

        stripped = line.strip()

        # ── Headings ─────────────────────────────────────────────────────────
        m = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            # Strip anchor links like {#section-id}
            text = re.sub(r"\{#[^}]+\}", "", text).strip()
            style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
            para = doc.add_heading(text, level=level)
            para.style = doc.styles[style_map.get(level, "Heading 3")]
            if level == 1:
                para.paragraph_format.space_before = Pt(18)
            i += 1
            continue

        # ── Horizontal rule ───────────────────────────────────────────────────
        if re.match(r"^---+$", stripped):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Bullet / checkbox lists ───────────────────────────────────────────
        m = re.match(r"^(\s*)[-*]\s+\[[ x]\]\s+(.*)", line)
        if m:
            indent, text = m.group(1), m.group(2)
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent = Inches(0.3 * (len(indent) // 2 + 1))
            apply_inline_code(para, text)
            i += 1
            continue

        m = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if m:
            indent, text = m.group(1), m.group(2)
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent = Inches(0.3 * (len(indent) // 2 + 1))
            apply_inline_code(para, text)
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if m:
            indent, text = m.group(1), m.group(2)
            para = doc.add_paragraph(style="List Number")
            para.paragraph_format.left_indent = Inches(0.3 * (len(indent) // 2 + 1))
            apply_inline_code(para, text)
            i += 1
            continue

        # ── Blank line ────────────────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── Normal paragraph ──────────────────────────────────────────────────
        # Strip markdown bold/italic markers, keep inline code
        text = stripped
        text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
        text = re.sub(r"\*([^*]+)\*", r"\1", text)
        text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)  # links

        para = doc.add_paragraph()
        apply_inline_code(para, text)
        i += 1

    flush_code()
    flush_table()

    doc.save(out_path)
    print(f"Written: {out_path}")


if __name__ == "__main__":
    root = Path(__file__).parent
    md_path = root / "DEPLOYMENT.md"
    out_path = root / "DEPLOYMENT.docx"

    if not md_path.exists():
        sys.exit(f"Not found: {md_path}")

    build_docx(md_path, out_path)
